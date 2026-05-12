from __future__ import annotations

import importlib.util
import json
import sys
from pathlib import Path

from docx import Document


ROOT_DIR = Path(__file__).resolve().parents[1]
SCRIPT_PATH = ROOT_DIR / "scripts" / "01_parse_docx.py"
SOURCE_DOC_PATH = ROOT_DIR / "data" / "raw" / "All database - NEO CITY.docx"


def load_parse_docx_module():
    spec = importlib.util.spec_from_file_location("parse_docx_script", SCRIPT_PATH)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Unable to load parser module from {SCRIPT_PATH}")

    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


parse_docx = load_parse_docx_module()


def build_sample_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("FACTSHEET", style="Title")
    document.add_paragraph("Đây là đoạn giới thiệu dự án.")

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Hạng mục"
    table.cell(0, 1).text = "Giá trị"
    table.cell(1, 0).text = "Quy mô"
    table.cell(1, 1).text = "60ha"

    document.add_paragraph("LIÊN KẾT VÙNG DỰ ÁN", style="Title")
    document.add_paragraph("Kết nối giao thông thuận tiện.")
    document.save(path)


def test_known_heading_detection() -> None:
    expected_matches = {
        "FACTSHEET": "factsheet",
        "LIÊN KẾT VÙNG DỰ ÁN": "location_connectivity",
        "Personas TA": "personas",
        "Concept": "concept_positioning",
        "Định vị": "concept_positioning",
        "CHIẾN LƯỢC BÁN HÀNG": "sales_strategy",
        "CHÍNH SÁCH BÁN HÀNG": "sales_policy",
        "Tình trạng pháp lý": "legal",
        "Giá bán & Chính sách bán hàng": "pricing",
        "Bản thị trường khu vực Mê Linh": "market",
        "Phiếu tính giá theo từng CSBH": "price_sheet",
    }

    for heading, expected_section_key in expected_matches.items():
        assert parse_docx.match_section_key(heading) == expected_section_key


def test_section_start_detection_uses_major_heading_styles() -> None:
    assert parse_docx.detect_section_start("Định vị", "Title") == "concept_positioning"
    assert parse_docx.detect_section_start("Định vị", "Heading 3") is None
    assert parse_docx.detect_section_start("CHÍNH SÁCH BÁN HÀNG", "Heading 1") == (
        "sales_policy"
    )


def test_output_file_creation_and_table_preservation(tmp_path: Path) -> None:
    document_path = tmp_path / "sample.docx"
    output_path = tmp_path / "neo_city_sections.json"
    build_sample_docx(document_path)

    sections = parse_docx.parse_docx_to_json(document_path, output_path)

    assert output_path.exists()
    payload = json.loads(output_path.read_text(encoding="utf-8"))
    assert len(sections) == 2
    assert len(payload) == 2
    assert payload[0]["section_key"] == "factsheet"
    assert "| Hạng mục | Giá trị |" in payload[0]["raw_text"]
    assert "| --- | --- |" in payload[0]["raw_text"]
    assert "| Quy mô | 60ha |" in payload[0]["raw_text"]


def test_actual_document_sections_have_non_empty_raw_text() -> None:
    sections = parse_docx.extract_sections(SOURCE_DOC_PATH)

    assert len(sections) == len(parse_docx.SECTION_TITLES)
    assert [section.source_title for section in sections] == [
        title for title, _ in parse_docx.SECTION_TITLES
    ]
    assert all(section.raw_text.strip() for section in sections)
